"""
OCR processing using Azure AI Document Intelligence.
Extracts text and tables from PDFs and Word documents while preserving structure.
"""

from typing import Dict, Any, List, Optional
import os
import json
import base64

from config import get_settings

settings = get_settings()

# Track if we've already logged OCR errors (avoid spam)
_ocr_error_logged = False

# Content type mapping based on file extension
CONTENT_TYPE_MAP = {
    ".pdf": "application/pdf",
    ".doc": "application/msword",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def get_content_type(file_path: str) -> str:
    """Get the content type based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()
    return CONTENT_TYPE_MAP.get(ext, "application/pdf")


async def process_document_ocr(file_path: str) -> Dict[str, Any]:
    """
    Process a document using Azure Document Intelligence.
    
    Uses the "prebuilt-layout" model which:
    - Extracts text with reading order
    - Preserves document structure (headers, paragraphs)
    - Extracts tables as structured data
    - Returns page-by-page results
    
    Supports PDF and Word documents (.doc, .docx).
    """
    global _ocr_error_logged
    
    # Check if Azure Document Intelligence is configured
    if not settings.azure_document_intelligence_endpoint or not settings.azure_document_intelligence_key:
        if not _ocr_error_logged:
            print("[OCR] Azure Document Intelligence not configured. Using fallback extraction.")
            _ocr_error_logged = True
        return await mock_ocr_result(file_path)
    
    try:
        from azure.ai.documentintelligence import DocumentIntelligenceClient
        from azure.ai.documentintelligence.models import AnalyzeDocumentRequest
        from azure.core.credentials import AzureKeyCredential
        
        # Validate endpoint format
        endpoint = settings.azure_document_intelligence_endpoint.rstrip('/')
        if not endpoint.startswith('https://'):
            endpoint = f"https://{endpoint}"
        
        # Ensure endpoint ends correctly (Azure DI expects base URL without path)
        if '/documentintelligence' in endpoint.lower():
            endpoint = endpoint.split('/documentintelligence')[0]
        
        # Initialize client
        client = DocumentIntelligenceClient(
            endpoint=endpoint,
            credential=AzureKeyCredential(settings.azure_document_intelligence_key)
        )
        
        # Read file and encode to base64
        with open(file_path, "rb") as f:
            document_bytes = f.read()
        
        bytes_source = base64.b64encode(document_bytes).decode('utf-8')
        
        # Create request with bytes_source (SDK 1.0.x format)
        analyze_request = AnalyzeDocumentRequest(bytes_source=bytes_source)
        
        # Analyze document using SDK 1.0.x API
        poller = client.begin_analyze_document(
            model_id="prebuilt-layout",
            body=analyze_request
        )
        
        result = poller.result()
        
        # Process result
        return process_azure_result(result)
        
    except ImportError as e:
        if not _ocr_error_logged:
            print(f"[OCR] Azure Document Intelligence SDK not installed: {e}")
            _ocr_error_logged = True
        return await mock_ocr_result(file_path)
    except Exception as e:
        error_str = str(e)
        if not _ocr_error_logged:
            print(f"[OCR] Azure Document Intelligence error: {error_str}")
            # Provide helpful debugging info
            if "404" in error_str:
                print(f"[OCR] 404 Error - Check your endpoint URL: {settings.azure_document_intelligence_endpoint}")
                print("[OCR] Endpoint should be: https://<resource-name>.cognitiveservices.azure.com")
            elif "401" in error_str or "403" in error_str:
                print("[OCR] Authentication error - Check your API key")
            _ocr_error_logged = True
        return await mock_ocr_result(file_path)


def process_azure_result(result) -> Dict[str, Any]:
    """Process Azure Document Intelligence result into structured format."""
    
    pages = []
    tables = []
    
    # Extract pages
    for page_num, page in enumerate(result.pages, 1):
        page_text = ""
        paragraphs = []
        
        # Get paragraphs for this page
        if result.paragraphs:
            for para in result.paragraphs:
                if para.bounding_regions and para.bounding_regions[0].page_number == page_num:
                    paragraphs.append({
                        "text": para.content,
                        "role": para.role if hasattr(para, "role") else None
                    })
                    page_text += para.content + "\n"
        
        pages.append({
            "page_number": page_num,
            "text": page_text,
            "paragraphs": paragraphs,
            "width": page.width,
            "height": page.height
        })
    
    # Extract tables
    if result.tables:
        for table_num, table in enumerate(result.tables, 1):
            cells = []
            for cell in table.cells:
                cells.append({
                    "row": cell.row_index,
                    "col": cell.column_index,
                    "text": cell.content,
                    "is_header": cell.kind == "columnHeader" if hasattr(cell, "kind") else False
                })
            
            # Get page number
            page_num = table.bounding_regions[0].page_number if table.bounding_regions else 1
            
            tables.append({
                "table_number": table_num,
                "page": page_num,
                "rows": table.row_count,
                "cols": table.column_count,
                "cells": cells
            })
    
    return {
        "page_count": len(pages),
        "pages": pages,
        "tables": tables,
        "full_text": "\n\n".join(p["text"] for p in pages)
    }


async def mock_ocr_result(file_path: str) -> Dict[str, Any]:
    """
    Generate mock OCR result for demo/development.
    Attempts to extract text using pypdf for PDFs or python-docx for Word documents.
    """
    
    ext = os.path.splitext(file_path)[1].lower()
    
    # Handle Word documents
    if ext in ('.doc', '.docx'):
        return await mock_ocr_word(file_path)
    
    # Handle PDFs
    try:
        from pypdf import PdfReader
        
        reader = PdfReader(file_path)
        pages = []
        
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            pages.append({
                "page_number": page_num,
                "text": text,
                "paragraphs": [{"text": p, "role": None} for p in text.split("\n\n") if p.strip()]
            })
        
        return {
            "page_count": len(pages),
            "pages": pages,
            "tables": [],  # pypdf doesn't extract tables well
            "full_text": "\n\n".join(p["text"] for p in pages),
            "_mock": True
        }
        
    except ImportError:
        # No pypdf, return basic mock
        return {
            "page_count": 10,
            "pages": [
                {"page_number": i, "text": f"[Mock page {i} content]", "paragraphs": []}
                for i in range(1, 11)
            ],
            "tables": [],
            "full_text": "[Mock document content - pypdf not installed]",
            "_mock": True
        }
    except Exception as e:
        print(f"Mock OCR failed: {e}")
        return {
            "page_count": 0,
            "pages": [],
            "tables": [],
            "full_text": "",
            "error": str(e)
        }


async def mock_ocr_word(file_path: str) -> Dict[str, Any]:
    """
    Extract text from Word documents for mock/demo purposes.
    Uses python-docx for .docx files.
    """
    try:
        from docx import Document as DocxDocument
        
        doc = DocxDocument(file_path)
        paragraphs = []
        full_text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    "text": para.text,
                    "role": "title" if para.style.name.startswith("Heading") else None
                })
                full_text_parts.append(para.text)
        
        # Extract tables
        tables = []
        for table_num, table in enumerate(doc.tables, 1):
            cells = []
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cells.append({
                        "row": row_idx,
                        "col": col_idx,
                        "text": cell.text,
                        "is_header": row_idx == 0
                    })
            tables.append({
                "table_number": table_num,
                "page": 1,  # Word docs don't have pages in the same way
                "rows": len(table.rows),
                "cols": len(table.columns),
                "cells": cells
            })
        
        # Word docs are treated as single page for mock purposes
        return {
            "page_count": 1,
            "pages": [{
                "page_number": 1,
                "text": "\n\n".join(full_text_parts),
                "paragraphs": paragraphs
            }],
            "tables": tables,
            "full_text": "\n\n".join(full_text_parts),
            "_mock": True
        }
        
    except ImportError:
        return {
            "page_count": 1,
            "pages": [{"page_number": 1, "text": "[Mock Word content - python-docx not installed]", "paragraphs": []}],
            "tables": [],
            "full_text": "[Mock Word document content - python-docx not installed]",
            "_mock": True
        }
    except Exception as e:
        print(f"Word OCR failed: {e}")
        return {
            "page_count": 0,
            "pages": [],
            "tables": [],
            "full_text": "",
            "error": str(e)
        }


def extract_tables_as_text(tables: List[Dict[str, Any]]) -> str:
    """Convert extracted tables to readable text format."""
    
    result = []
    
    for table in tables:
        rows_data = {}
        
        # Group cells by row
        for cell in table.get("cells", []):
            row = cell["row"]
            if row not in rows_data:
                rows_data[row] = {}
            rows_data[row][cell["col"]] = cell["text"]
        
        # Format as text table
        table_text = f"Table {table['table_number']} (Page {table['page']}):\n"
        
        for row_idx in sorted(rows_data.keys()):
            row = rows_data[row_idx]
            row_text = " | ".join(
                row.get(col, "") for col in range(table.get("cols", 0))
            )
            table_text += row_text + "\n"
        
        result.append(table_text)
    
    return "\n\n".join(result)




